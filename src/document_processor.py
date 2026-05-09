from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List

from src.table_structure import extract_structure_tables


## Document processor: PDF and image inputs use PaddleOCR 3.x inside the paddleocr3 container.
class DocumentProcessor:
    ## Keep config and cache the OCR engine so batch OCR does not reload models for every file.
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self._ocr_engine = None
        self._structure_engine = None

    ## Dispatch by file extension. PDF and images go directly to PaddleOCR.
    def process(self, path: Path) -> Dict[str, Any]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._process_pdf_by_paddleocr(path)
        if suffix == ".docx":
            return self._process_docx(path)
        if suffix == ".json":
            return self._process_json(path)
        if suffix in {".md", ".markdown", ".txt"}:
            return self._process_text(path)
        if suffix in {".png", ".jpg", ".jpeg"}:
            return self._process_image_by_paddleocr(path)
        raise ValueError(f"Unsupported input type: {path}")

    ## Create the shared document structure used by downstream RAG/model/Excel stages.
    def _base_document(self, path: Path, kind: str) -> Dict[str, Any]:
        return {
            "source_file": str(path),
            "file_name": path.name,
            "type": kind,
            "pages": [],
            "metadata": {},
        }

    ## Process PDF by passing the PDF path directly to PaddleOCR in the paddleocr3 container.
    def _process_pdf_by_paddleocr(self, path: Path) -> Dict[str, Any]:
        document = self._base_document(path, "pdf")
        ocr_result = self._ocr_file(path)
        pages = ocr_result.get("pages", [])
        tables_by_page = ocr_result.get("tables_by_page", {})
        document["metadata"]["ocr_engine"] = "paddleocr3"
        document["metadata"]["ocr_status"] = ocr_result.get("status", "")
        document["metadata"]["ocr_artifacts"] = ocr_result.get("artifacts", {})
        document["metadata"]["table_structure_status"] = ocr_result.get("table_structure_status", "")
        document["metadata"]["table_structure_artifacts"] = ocr_result.get("table_structure_artifacts", {})

        for page in pages:
            page_number = page.get("page", 1)
            document["pages"].append(
                {
                    "page": page_number,
                    "text": page.get("text", ""),
                    "tables": tables_by_page.get(int(page_number), []),
                    "images": [],
                    "raw_ocr": page if self.config.get("ocr", {}).get("save_raw_ocr", True) else {},
                    "text_source": "paddleocr3",
                }
            )

        if not document["pages"]:
            document["pages"].append(
                {
                    "page": 1,
                    "text": "",
                    "tables": [],
                    "images": [],
                    "raw_ocr": ocr_result,
                    "text_source": "paddleocr3",
                }
            )
        return document

    ## Process image by passing the image path directly to PaddleOCR.
    def _process_image_by_paddleocr(self, path: Path) -> Dict[str, Any]:
        document = self._base_document(path, "image")
        ocr_result = self._ocr_file(path)
        first_page = ocr_result.get("pages", [{}])[0] if ocr_result.get("pages") else {}
        tables_by_page = ocr_result.get("tables_by_page", {})
        document["metadata"]["ocr_engine"] = "paddleocr3"
        document["metadata"]["ocr_status"] = ocr_result.get("status", "")
        document["metadata"]["ocr_artifacts"] = ocr_result.get("artifacts", {})
        document["metadata"]["table_structure_status"] = ocr_result.get("table_structure_status", "")
        document["metadata"]["table_structure_artifacts"] = ocr_result.get("table_structure_artifacts", {})
        document["pages"].append(
            {
                "page": 1,
                "text": first_page.get("text", ""),
                "tables": tables_by_page.get(1, []),
                "images": [str(path)],
                "raw_ocr": first_page if self.config.get("ocr", {}).get("save_raw_ocr", True) else {},
                "text_source": "paddleocr3",
            }
        )
        return document

    ## Process Word documents without OCR, keeping paragraph text and tables.
    def _process_docx(self, path: Path) -> Dict[str, Any]:
        from docx import Document

        docx_file = Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in docx_file.paragraphs if paragraph.text.strip()]
        tables = []
        for table_index, table in enumerate(docx_file.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append({"table_index": table_index, "rows": rows})

        document = self._base_document(path, "docx")
        document["pages"].append(
            {"page": 1, "text": "\n".join(paragraphs), "tables": tables, "images": [], "raw_ocr": {}, "text_source": "docx"}
        )
        return document

    ## Process JSON by preserving the raw object and exposing it as readable text.
    def _process_json(self, path: Path) -> Dict[str, Any]:
        data = json.loads(path.read_text(encoding="utf-8"))
        document = self._base_document(path, "json")
        text = json.dumps(data, ensure_ascii=False, indent=2)
        document["pages"].append({"page": 1, "text": text, "tables": [], "images": [], "raw_ocr": data, "text_source": "json"})
        return document

    ## Process Markdown/text by reading it directly.
    def _process_text(self, path: Path) -> Dict[str, Any]:
        document = self._base_document(path, path.suffix.lower().lstrip("."))
        document["pages"].append(
            {"page": 1, "text": path.read_text(encoding="utf-8", errors="ignore"), "tables": [], "images": [], "raw_ocr": {}, "text_source": "text"}
        )
        return document

    ## Run PaddleOCR on a PDF or image path.
    def _ocr_file(self, path: Path) -> Dict[str, Any]:
        if not self.config.get("ocr", {}).get("enabled", True):
            return {"status": "ocr_disabled", "pages": []}

        ocr = self._get_ocr_engine()
        if ocr is None:
            return {"status": "paddleocr3_unavailable", "pages": []}

        raw_result = run_paddleocr(ocr, path)
        artifacts = save_paddleocr_artifacts(raw_result, path, self.config)
        pages = normalize_paddleocr_result(raw_result)
        structure_status = "disabled"
        structure_artifacts: Dict[str, Any] = {}
        tables_by_page: Dict[int, List[Dict[str, Any]]] = {}
        if self.config.get("ocr", {}).get("table_structure", False):
            structure = self._get_structure_engine()
            if structure is None:
                structure_status = "pp_structure_v3_unavailable"
            else:
                try:
                    structure_result = run_structure_recognition(structure, path)
                    structure_artifacts = save_structure_artifacts(structure_result, path, self.config)
                    tables_by_page = extract_structure_tables(structure_result, pages)
                    structure_status = "ok"
                except Exception as exc:
                    structure_status = f"error: {exc}"

        return {
            "status": "ok",
            "pages": pages,
            "artifacts": artifacts,
            "table_structure_status": structure_status,
            "table_structure_artifacts": structure_artifacts,
            "tables_by_page": tables_by_page,
        }

    ## Initialize PaddleOCR lazily inside the paddleocr3 container.
    def _get_ocr_engine(self) -> Any:
        if self._ocr_engine is not None:
            return self._ocr_engine
        try:
            from paddleocr import PaddleOCR
        except Exception:
            return None
        self._ocr_engine = build_paddleocr_engine(PaddleOCR, self.config)
        return self._ocr_engine

    ## Initialize the document structure engine lazily. This is optional because some containers only ship OCR.
    def _get_structure_engine(self) -> Any:
        if self._structure_engine is not None:
            return self._structure_engine
        try:
            from paddleocr import PPStructureV3
        except Exception:
            return None
        try:
            self._structure_engine = build_structure_engine(PPStructureV3, self.config)
        except Exception:
            return None
        return self._structure_engine



## Normalize PaddleOCR 3.x output to [{page, text, items}], preserving box/text/score per OCR line.
def normalize_paddleocr_result(raw_result: Any) -> List[Dict[str, Any]]:
    pages: List[Dict[str, Any]] = []
    if raw_result is None:
        return pages

    for page_index, page_result in enumerate(raw_result or [], start=1):
        page_data = paddle_result_to_dict(page_result)
        items = extract_ocr3_items(page_data)
        page_text = "\n".join(item["text"] for item in items if item.get("text"))
        pages.append(
            {
                "page": int(page_data.get("page_index", page_index - 1)) + 1 if "page_index" in page_data else page_index,
                "text": page_text,
                "items": items,
                "raw_keys": sorted(str(key) for key in page_data.keys()),
            }
        )

    return pages


## Convert PaddleOCR 3.x OCRResult/dict output into a regular dict.
def paddle_result_to_dict(page_result: Any) -> Dict[str, Any]:
    if isinstance(page_result, dict):
        return page_result

    json_value = getattr(page_result, "json", None)
    if isinstance(json_value, dict):
        return json_value
    if callable(json_value):
        value = json_value()
        if isinstance(value, dict):
            return value

    data = getattr(page_result, "data", None)
    if isinstance(data, dict):
        return data

    result = {}
    for name in ("rec_texts", "rec_scores", "rec_polys", "rec_boxes", "dt_polys", "page_index"):
        if hasattr(page_result, name):
            result[name] = getattr(page_result, name)
    return result


## Extract OCR line items from PaddleOCR 3.x rec_texts/rec_scores/rec_polys fields.
def extract_ocr3_items(page_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    rec_texts = page_data.get("rec_texts") or page_data.get("texts") or page_data.get("text")
    if isinstance(rec_texts, str):
        rec_texts = [rec_texts]
    if not rec_texts:
        return []

    rec_scores = page_data.get("rec_scores") or page_data.get("scores") or []
    rec_polys = page_data.get("rec_polys") or page_data.get("rec_boxes") or page_data.get("dt_polys") or []

    items: List[Dict[str, Any]] = []
    for index, text in enumerate(rec_texts):
        score = rec_scores[index] if index < len(rec_scores) else None
        box = rec_polys[index] if index < len(rec_polys) else []
        items.append({"box": to_jsonable(box), "text": str(text), "score": to_jsonable(score)})
    return items


## Convert numpy arrays/scalars and other non-JSON-native values to JSON-friendly data.
def to_jsonable(value: Any) -> Any:
    if hasattr(value, "tolist"):
        return value.tolist()
    if hasattr(value, "item"):
        return value.item()
    return value


## Save PaddleOCR 3.x visualization images and raw result JSON next to OCR text output.
def save_paddleocr_artifacts(raw_result: Any, source_path: Path, config: Dict[str, Any]) -> Dict[str, Any]:
    ocr_config = config.get("ocr", {})
    if not ocr_config.get("save_visualization", True):
        return {}

    output_root = Path(config.get("input", {}).get("output_dir", "output"))
    run_dir = output_root / source_path.stem
    visual_dir = run_dir / ocr_config.get("visualization_dir_name", "ocr_visualization")
    json_dir = run_dir / ocr_config.get("result_json_dir_name", "ocr_result_json")
    visual_dir.mkdir(parents=True, exist_ok=True)
    json_dir.mkdir(parents=True, exist_ok=True)

    for page_result in raw_result or []:
        if hasattr(page_result, "save_to_img"):
            page_result.save_to_img(str(visual_dir))
        if hasattr(page_result, "save_to_json"):
            page_result.save_to_json(str(json_dir), indent=2, ensure_ascii=False)

    return {
        "visualization_dir": str(visual_dir),
        "result_json_dir": str(json_dir),
        "visualization_files": sorted(str(path) for path in visual_dir.glob("*")),
        "result_json_files": sorted(str(path) for path in json_dir.glob("*")),
    }


## Save PP-StructureV3 artifacts in separate folders so text OCR and layout/table results are easy to compare.
def save_structure_artifacts(raw_result: Any, source_path: Path, config: Dict[str, Any]) -> Dict[str, Any]:
    ocr_config = config.get("ocr", {})
    if not ocr_config.get("save_visualization", True):
        return {}

    output_root = Path(config.get("input", {}).get("output_dir", "output"))
    run_dir = output_root / source_path.stem
    structure_dir = run_dir / ocr_config.get("structure_result_dir_name", "structure_result")
    structure_dir.mkdir(parents=True, exist_ok=True)

    for page_result in raw_result or []:
        if hasattr(page_result, "save_to_json"):
            page_result.save_to_json(str(structure_dir), indent=2, ensure_ascii=False)
        if hasattr(page_result, "save_to_markdown"):
            page_result.save_to_markdown(str(structure_dir))
        if hasattr(page_result, "save_to_html"):
            page_result.save_to_html(str(structure_dir))
        if hasattr(page_result, "save_to_xlsx"):
            page_result.save_to_xlsx(str(structure_dir))

    return {
        "structure_result_dir": str(structure_dir),
        "structure_result_files": sorted(str(path) for path in structure_dir.glob("*")),
    }


## Build PaddleOCR 3.x engine only for the paddleocr3 container.
def build_paddleocr_engine(paddleocr_cls: Any, config: Dict[str, Any]) -> Any:
    ocr_config = config.get("ocr", {})
    return paddleocr_cls(
        lang=ocr_config.get("language", "ch"),
        device=ocr_config.get("device", "gpu:0"),
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_textline_orientation=True,
    )


## Build PP-StructureV3 for layout and table structure recognition.
def build_structure_engine(structure_cls: Any, config: Dict[str, Any]) -> Any:
    ocr_config = config.get("ocr", {})
    kwargs = {
        "lang": ocr_config.get("language", "ch"),
        "device": ocr_config.get("device", "gpu:0"),
        "use_doc_orientation_classify": False,
        "use_doc_unwarping": False,
        "use_textline_orientation": True,
        "use_table_recognition": True,
        "use_formula_recognition": ocr_config.get("use_formula_recognition", False),
        "use_chart_recognition": ocr_config.get("use_chart_recognition", False),
        "use_seal_recognition": ocr_config.get("use_seal_recognition", False),
    }
    try:
        return structure_cls(**kwargs)
    except TypeError:
        return structure_cls(
            device=ocr_config.get("device", "gpu:0"),
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=True,
        )


## Run PaddleOCR 3.x. The new container must provide predict().
def run_paddleocr(ocr: Any, path: Path) -> Any:
    return ocr.predict(str(path))


## Run PP-StructureV3. The result is parsed into tables and also saved as raw artifacts when possible.
def run_structure_recognition(structure: Any, path: Path) -> Any:
    return structure.predict(str(path))

